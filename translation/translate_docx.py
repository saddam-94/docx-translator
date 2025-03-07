import logging
import os
from docx import Document
from tqdm import tqdm 
import time
from multiprocessing import Pool, cpu_count
from functools import partial
import argostranslate.package
import argostranslate.translate
import random 
import yaml
# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

from .utils import preserve_special_elements, parse_args

def translate(text: str, source_lang: str, target_lang: str, verbose=False) -> str:
    """
    Translate text using argostranslate.

    Args:
        text (str): The text to translate
        source_lang (str): The source language of the text
        target_lang (str): The target language to translate to

    Returns:
        str: The translated text

    Note:
        - Uses MODEL_NAME and MODEL_URL from environment variables
        - Logs translations to 'translations.txt'
    """
    # TODO (ex.2) : You can modify this function except line 41-42

    # Download and install Argos Translate package
    argostranslate.package.update_package_index()
    available_packages = argostranslate.package.get_available_packages()
    package_to_install = next(
        filter(
            lambda x: x.from_code == source_lang and x.to_code == target_lang, available_packages
        )
    )
    argostranslate.package.install_from_path(package_to_install.download())

    ############ DO NOT CHANGE THIS BLOCK OF CODE ##############################
    # Simulate HTTP request latency with a random delay between 0.1 and 0.4 seconds
    # delay = random.uniform(0.1, 0.4)
    # time.sleep(delay)
    #############################################################################
    
    # Translate
    response = argostranslate.translate.translate(text, source_lang, target_lang)
    
    
    if verbose : 
        with open('translations.txt', 'a') as f:
            f.write(text + '=>'+ response + '\n')
    return response


def translate_runs_in_paragraph(paragraph: any, translation_func: callable, mode: str = "naive") -> None:
    """
    TODO (ex.1) : finish this doc
    Translates the text within a given paragraph, preserving special elements such as images, tables, 
    or other non-text content.

    Args:
        paragraph (any): The paragraph object (from libraries like `python-docx` or similar), which contains 
                          the text to be translated.
        translation_func (callable): The function to perform the translation, which should take a text string 
                                      as input and return the translated text.
        mode (str, optional): The translation mode. Default is "naive". In "naive" mode, the function simply 
                              translates the text runs. More advanced modes can be added if needed for 
                              specialized translation processes (e.g., handling graphs or tables differently).
                              also it offers Context-Aware Mode.

    Returns:
        None: This function modifies the paragraph in place by translating its runs and restoring any special 
              elements at their original positions.
    """
    # First preserve special elements with their positions
    special_elements = preserve_special_elements(paragraph)

    if mode == "naive":
        for run in paragraph.runs:
            if run.text.strip() and len(run.text) < 2000: # dodgy way to avoid translating long tables/graphs 
                translated_text = translation_func(run.text)
                run.text = translated_text.strip()+" "

    # Restore special elements in their original positions
    for original_pos, element in special_elements:
        # Insert the special element back at its relative position
        paragraph._p.insert(original_pos, element)

                

def translate_docx(args, input_path, output_path, source_lang, target_lang):
    """
    Translates the given DOCX document from the source language to the target language and saves it.

    Args:
        input_path (str): The path to the DOCX file to be translated.
        output_path (str): The path to save the translated DOCX file.
        source_lang (str): The source language code (e.g., 'en' for English).
        target_lang (str): The target language code (e.g., 'fr' for French).

    Returns:
        None: The function saves the translated document to the provided output path.

    Raises:
        ValueError: If the provided language pair is not supported or the translation package is not found.
        Exception: For any unexpected errors that occur during the translation process.
    """
    try:
        logger.info(f"Starting translation from {args.src_lang} to {args.tar_lang}")
        start = time.time()

        translation_func  = partial(translate, source_lang=source_lang, target_lang=target_lang)

        logger.info(f"Opening document: {input_path}")
        doc = Document(input_path)
        
        print(f"Translation in progress...")
        

        for paragraph in tqdm(doc.paragraphs):
            translate_runs_in_paragraph(paragraph, translation_func)

        for table in tqdm(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        translate_runs_in_paragraph(paragraph, translation_func)


        for section in tqdm(doc.sections):
            for paragraph in section.header.paragraphs:
                translate_runs_in_paragraph(paragraph, translation_func)
            for paragraph in  section.footer.paragraphs:
                translate_runs_in_paragraph(paragraph, translation_func)
            for table in  section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in tqdm(cell.paragraphs):
                            translate_runs_in_paragraph(paragraph, translation_func)

        # print(f"Translation is done !")

        # print(f"Saving translated document to: {output_path}")
        # doc.save(output_path)
        # end = time.time()
        # print(f"Total time taken: {end - start} seconds.")

        # Save the translated document
        logger.info(f"Saving translated document to: {output_path}")
        if not os.path.exists(output_path):
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        end = time.time()
        logger.info(f"Translation completed in {end - start:.2f} seconds.")
        return doc
    except FileNotFoundError as e:
        logger.error(f"Error: The input file was not found at {input_path}. Please check the path.")
        raise
    except ValueError as e:
        logger.error(f"Error: {e}")
        raise
    except Exception as e:
        logger.error(f"Unexpected error occurred: {str(e)}")
        raise
